from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re


# ── PDF ──────────────────────────────────────────────────────────


@dataclass(frozen=True)
class PDFLogicalPage:
    logical_page_number: int
    physical_page_number: int
    crop_box: tuple[float, float, float, float]
    text: str
    image_png: bytes
    embedded_image_count: int


_PAGE_MARKER = re.compile(r"第\s*(\d+)\s*页\s*[，,]\s*共\s*(\d+)\s*页")


def _should_split_pdf_page(
    *,
    width: float,
    height: float,
    left_text: str,
    right_text: str,
) -> bool:
    if width / max(height, 1) < 1.30:
        return False
    left_marker = _PAGE_MARKER.search(left_text)
    right_marker = _PAGE_MARKER.search(right_text)
    return bool(
        left_marker
        and right_marker
        and left_marker.group(1) != right_marker.group(1)
    )


def _render_pdf_crop(
    page: object,
    crop_box: tuple[float, float, float, float],
    resolution: int,
) -> bytes:
    cropped = page.crop(crop_box)
    image = cropped.to_image(resolution=resolution, antialias=True).original
    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def extract_pdf_pages(path: str, *, resolution: int = 144) -> list[PDFLogicalPage]:
    """Return ordered logical pages with crop text and rendered PNG bytes."""
    try:
        import pdfplumber
    except ModuleNotFoundError as exc:
        raise RuntimeError("需要安装 pdfplumber 来提取 PDF 页面") from exc

    logical_pages: list[PDFLogicalPage] = []
    with pdfplumber.open(path) as pdf:
        for physical_index, page in enumerate(pdf.pages, start=1):
            midpoint = page.width / 2
            left_box = (0.0, 0.0, midpoint, page.height)
            right_box = (midpoint, 0.0, page.width, page.height)
            left_text = _clean_extracted_text(page.crop(left_box).extract_text() or "")
            right_text = _clean_extracted_text(page.crop(right_box).extract_text() or "")
            if _should_split_pdf_page(
                width=page.width,
                height=page.height,
                left_text=left_text,
                right_text=right_text,
            ):
                crop_boxes = (left_box, right_box)
            else:
                crop_boxes = ((0.0, 0.0, page.width, page.height),)

            for crop_box in crop_boxes:
                cropped = page.crop(crop_box)
                text = _clean_extracted_text(cropped.extract_text() or "")
                marker = _PAGE_MARKER.search(text)
                logical_number = (
                    int(marker.group(1)) if marker else len(logical_pages) + 1
                )
                embedded_image_count = sum(
                    1
                    for image in page.images
                    if crop_box[0] <= float(image.get("x0", -1)) < crop_box[2]
                )
                logical_pages.append(
                    PDFLogicalPage(
                        logical_page_number=logical_number,
                        physical_page_number=physical_index,
                        crop_box=crop_box,
                        text=text,
                        image_png=_render_pdf_crop(page, crop_box, resolution),
                        embedded_image_count=embedded_image_count,
                    )
                )

    return sorted(logical_pages, key=lambda item: item.logical_page_number)

def extract_pdf_text(path: str) -> str:
    """提取 PDF 文本，优先用 pdfplumber（更准确），回退到 pypdf。"""
    text = _try_pdfplumber(path) or _try_pypdf(path)
    return _clean_extracted_text(text)


def _try_pdfplumber(path: str) -> str:
    try:
        import pdfplumber
    except ModuleNotFoundError:
        return ""
    try:
        pages_text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
        return "\n".join(pages_text).strip()
    except Exception:
        return ""


def _try_pypdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise RuntimeError("需要安装 pypdf 或 pdfplumber 来提取 PDF 文本") from exc

    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


# ── DOCX ─────────────────────────────────────────────────────────

def extract_docx_text(path: str) -> str:
    """提取 DOCX 文本，包括段落和表格。"""
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError("需要安装 python-docx 来提取 DOCX 文本") from exc

    document = Document(path)
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


# ── DOC (旧版 Word) ─────────────────────────────────────────────

def extract_doc_text(path: str) -> str:
    """提取 .doc 文本，尝试多种方式。"""
    text = _try_doc_with_docx(path) or _try_doc_binary_extract(path)
    if not text:
        raise RuntimeError(
            "无法提取 .doc 文件内容。请将文件转换为 .docx 或 .txt 格式后重新上传。"
        )
    return _clean_extracted_text(text)


def _try_doc_with_docx(path: str) -> str:
    """尝试用 python-docx 打开（某些 .doc 文件也兼容）。"""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _try_doc_binary_extract(path: str) -> str:
    """从 .doc 二进制中提取可读文本（简单方式）。"""
    try:
        data = Path(path).read_bytes()
        # 尝试 UTF-16LE（Word .doc 常用编码）
        text = data.decode("utf-16-le", errors="ignore")
        # 保留可打印字符：中文、英文、数字、常见标点
        cleaned = re.sub(
            r"[^一-鿿　-〿＀-￯a-zA-Z0-9\s.,;:!?()（）【】《》“”‘’、。，；：！？\-\n\r\t]",
            "", text
        )
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if len(cleaned) > 50:
            return cleaned
        # 回退：尝试直接提取连续中文
        chinese = re.findall(r"[一-鿿]{5,}", data.decode("latin-1", errors="ignore"))
        return "\n".join(chinese).strip() if chinese else ""
    except Exception:
        return ""


# ── TXT / MD ─────────────────────────────────────────────────────

def extract_plain_text(path: str) -> str:
    """提取纯文本或 Markdown 文件内容。"""
    data = Path(path).read_bytes()
    text = _decode_text(data)
    return _clean_text_content(text)


def _decode_text(data: bytes) -> str:
    """尝试多种编码解码。"""
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk", "gb2312", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _clean_text_content(text: str) -> str:
    """清理文本内容。"""
    # 移除 BOM
    text = text.lstrip("﻿")
    # 合并连续空行
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 清理每行首尾空格
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


# ── 通用清理 ─────────────────────────────────────────────────────

def _clean_extracted_text(text: str) -> str:
    """清理提取的文本：多余空格、断行等。"""
    if not text:
        return text
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


# ── 入口函数 ─────────────────────────────────────────────────────

def extract_text(path: str, mime_type: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    normalized_mime = (mime_type or "").lower()

    if suffix == ".pdf" or normalized_mime == "application/pdf":
        return extract_pdf_text(path)

    if suffix == ".docx" or normalized_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx_text(path)

    if suffix == ".doc" or normalized_mime == "application/msword":
        return extract_doc_text(path)

    if suffix == ".md" or normalized_mime in ("text/markdown", "text/x-markdown"):
        return extract_plain_text(path)

    if suffix == ".txt" or normalized_mime.startswith("text/") or not suffix:
        return extract_plain_text(path)

    # 未知格式，尝试当纯文本读取
    return extract_plain_text(path)
